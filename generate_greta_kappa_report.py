#!/usr/bin/env python3
"""Generate Greta's Kappa report and all-entry working translation document."""

from __future__ import annotations

import argparse
from collections import defaultdict
from datetime import date
import json
from pathlib import Path
import re
from statistics import mean
from typing import Any

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from db import get_connection
from generate_translation_prompt_evaluation import (
    TranslationMetricEvaluator,
    build_pair_rows,
    metric_range,
    normalize_text,
)
from paper_corpus import paper_kappa_review_cte_body
from source_documents import source_document_priority_sql


MAIN_PROFILE = "gpt-5.5"
MAIN_PROMPT_VERSION = 3
REPEAT_PROFILE = "gpt-5.5_v3_repeat"
REPEAT_PROFILE_VERSION = 1
EXPECTED_PAPER_PASSAGES = 100
EXPECTED_REPEAT_RUNS = 500
MODEL_TIMELINE_PROFILES = (
    "gpt-4-turbo-2024-04-09",
    "gpt-4o-2024-05-13",
    "gpt-4o-2024-08-06",
    "gpt-4o-2024-11-20",
    "gpt-4.1-2025-04-14",
    "gpt-5",
    "gpt-5.1",
    "gpt-5.2",
    "gpt-5.3-chat-latest",
    "gpt-5.4",
    "gpt-5.5",
    "gpt-5.6-sol",
)
MODEL_TIMELINE_PROMPT_VERSIONS = (1, 2, 3)
EXPECTED_MODEL_TIMELINE_PAIRS = (
    len(MODEL_TIMELINE_PROFILES)
    * len(MODEL_TIMELINE_PROMPT_VERSIONS)
    * EXPECTED_PAPER_PASSAGES
)

NAVY = "17324D"
BLUE = "2A658C"
PALE_BLUE = "EAF2F7"
PALE_GREEN = "E8F4EC"
PALE_ORANGE = "FFF1D6"
PALE_GREY = "F2F3F5"
DARK_GREY = "40464D"
WHITE = "FFFFFF"


def compact_whitespace(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def plain_markdown(value: Any) -> str:
    return re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)", r"\1", compact_whitespace(value))


def english_headword(translation: str, fallback: str) -> str:
    text = plain_markdown(translation)
    candidates = [position for position in (text.find(":"), text.find(",")) if position > 0]
    if candidates:
        text = text[: min(candidates)]
    text = text.strip(" \"'[]()")
    return text or fallback


def formatted_billerbeck(value: str, entry_number: int) -> str:
    value = compact_whitespace(value)
    match = re.fullmatch(r"([ΚK])\s*(\d+)", value)
    if match:
        return f"Κ {match.group(2)}"
    return value or f"Κ {entry_number}"


def entry_title(entry: dict[str, Any]) -> str:
    translation = entry.get("human_translation") or entry.get("v3_translation") or ""
    headword = english_headword(str(translation), str(entry.get("lemma") or "Untitled"))
    references = []
    billerbeck = formatted_billerbeck(
        str(entry.get("billerbeck_id") or ""),
        int(entry["entry_number"]),
    )
    if billerbeck:
        references.append(f"{billerbeck} Billerbeck")
    meineke = compact_whitespace(entry.get("meineke_id"))
    if meineke:
        references.append(f"{meineke} Meineke")
    reference_text = " = ".join(references)
    return f"{headword} ({reference_text})" if reference_text else headword


def json_safe(value: Any) -> Any:
    if isinstance(value, dict):
        return {key: json_safe(item) for key, item in value.items()}
    if isinstance(value, list):
        return [json_safe(item) for item in value]
    if isinstance(value, tuple):
        return [json_safe(item) for item in value]
    if hasattr(value, "isoformat"):
        return value.isoformat()
    return value


def fetch_entries(cur) -> list[dict[str, Any]]:
    query = f"""
        SELECT
            a.id AS lemma_id,
            a.entry_number,
            a.lemma,
            a.billerbeck_id,
            a.meineke_id,
            COALESCE(a.human_notes, '') AS human_notes,
            ht.id AS human_translation_id,
            COALESCE(ht.translation_text, '') AS human_translation,
            stv.id AS source_text_version_id,
            COALESCE(stv.source_document, '') AS source_document,
            tr.id AS v3_run_id,
            COALESCE(tr.translation_text, '') AS v3_translation
        FROM assembled_lemmas a
        LEFT JOIN LATERAL (
            SELECT candidate.*
            FROM human_translations candidate
            WHERE candidate.lemma_id = a.id
              AND candidate.status = 'approved'
              AND candidate.stage IN ('reviewed', 'final')
              AND NULLIF(BTRIM(candidate.translation_text), '') IS NOT NULL
            ORDER BY
                (candidate.stage = 'final') DESC,
                COALESCE(candidate.reviewed_at, candidate.updated_at, candidate.created_at) DESC,
                candidate.id DESC
            LIMIT 1
        ) ht ON TRUE
        LEFT JOIN LATERAL (
            SELECT candidate.id, candidate.source_document
            FROM lemma_source_text_versions candidate
            WHERE candidate.lemma_id = a.id
              AND candidate.is_current = TRUE
            ORDER BY {source_document_priority_sql("candidate.source_document")},
                     candidate.id DESC
            LIMIT 1
        ) stv ON TRUE
        LEFT JOIN LATERAL (
            SELECT candidate.*
            FROM translation_runs candidate
            JOIN translation_prompt_profiles p ON p.id = candidate.profile_id
            JOIN translation_prompt_profile_versions pv ON pv.id = candidate.profile_version_id
            WHERE candidate.lemma_id = a.id
              AND p.name = %s
              AND pv.version = %s
              AND candidate.status IN ('completed', 'approved')
              AND NULLIF(BTRIM(candidate.translation_text), '') IS NOT NULL
            ORDER BY
                CASE WHEN candidate.source_text_version_id = stv.id THEN 0 ELSE 1 END,
                CASE candidate.status WHEN 'approved' THEN 0 ELSE 1 END,
                COALESCE(candidate.reviewed_at, candidate.completed_at, candidate.created_at) DESC,
                candidate.id DESC
            LIMIT 1
        ) tr ON TRUE
        WHERE a.version = 'epitome'
          AND a.entry_number BETWEEN 1 AND 317
          AND LEFT(BTRIM(COALESCE(a.lemma, '')), 1) IN ('Κ', 'κ')
          AND NULLIF(BTRIM(a.billerbeck_id), '') IS NOT NULL
          AND NULLIF(BTRIM(a.meineke_id), '') IS NOT NULL
        ORDER BY a.entry_number, a.id
    """
    cur.execute(query, (MAIN_PROFILE, MAIN_PROMPT_VERSION))
    entries = [dict(row) for row in cur.fetchall()]
    if len(entries) != 317:
        raise RuntimeError(f"Expected 317 official Kappa entries, found {len(entries)}")
    return entries


def attach_notes_and_sources(cur, entries: list[dict[str, Any]]) -> None:
    lemma_ids = [int(entry["lemma_id"]) for entry in entries]
    cur.execute(
        """
        SELECT
            lemma_id,
            id,
            proper_noun,
            lemma_form,
            english_translation,
            citation,
            work_title,
            effective_wikidata_qid,
            effective_resolution_status,
            effective_resolution_source
        FROM effective_proper_nouns
        WHERE lemma_id = ANY(%s)
          AND role = 'source'
        ORDER BY lemma_id, id
        """,
        (lemma_ids,),
    )
    cited_sources: dict[int, list[dict[str, Any]]] = defaultdict(list)
    for row in cur.fetchall():
        cited_sources[int(row["lemma_id"])].append(dict(row))

    cur.execute(
        """
        SELECT
            lemma_id,
            id,
            phrase_text,
            commentary_text,
            note_kind,
            generation_source,
            review_status,
            publication_status,
            confidence
        FROM lemma_commentary_entries
        WHERE lemma_id = ANY(%s)
          AND stale_at IS NULL
          AND NULLIF(BTRIM(commentary_text), '') IS NOT NULL
          AND generation_source = 'human'
          AND review_status = 'approved'
          AND publication_status = 'public_reviewed'
        ORDER BY lemma_id, id
        """,
        (lemma_ids,),
    )
    commentary: dict[int, list[dict[str, Any]]] = defaultdict(list)
    for row in cur.fetchall():
        commentary[int(row["lemma_id"])].append(dict(row))

    for entry in entries:
        lemma_id = int(entry["lemma_id"])
        entry["cited_sources"] = cited_sources.get(lemma_id, [])
        entry["commentary_entries"] = commentary.get(lemma_id, [])
        entry["title"] = entry_title(entry)
        entry["translation_status"] = (
            "FINAL" if compact_whitespace(entry.get("human_translation")) else "UNCHECKED"
        )


def fetch_paper_entries(cur, entries: list[dict[str, Any]]) -> list[dict[str, Any]]:
    cur.execute(
        f"""
        WITH {paper_kappa_review_cte_body("paper_corpus")}
        SELECT
            pc.corpus_order,
            pc.corpus_source_row_id,
            pc.kappa_review_row_id,
            pc.lemma_id,
            pc.human_translation_id,
            ht.translation_text AS paper_human_translation
        FROM paper_corpus pc
        JOIN human_translations ht ON ht.id = pc.human_translation_id
        ORDER BY pc.corpus_order
        """
    )
    paper_rows = [dict(row) for row in cur.fetchall()]
    if len(paper_rows) != EXPECTED_PAPER_PASSAGES:
        raise RuntimeError(
            f"Expected {EXPECTED_PAPER_PASSAGES} paper passages, found {len(paper_rows)}"
        )
    entries_by_id = {int(entry["lemma_id"]): entry for entry in entries}
    result = []
    for paper_row in paper_rows:
        entry = dict(entries_by_id[int(paper_row["lemma_id"])])
        entry.update(paper_row)
        entry["human_translation"] = paper_row["paper_human_translation"]
        entry["title"] = entry_title(entry)
        result.append(entry)
    return result


def fetch_repeatability_rows(cur) -> list[dict[str, Any]]:
    cur.execute(
        f"""
        WITH {paper_kappa_review_cte_body("paper_corpus")}
        SELECT
            pc.corpus_order,
            tr.id AS run_id,
            tr.lemma_id,
            tr.run_index,
            tr.model,
            tr.temperature,
            tr.top_p,
            tr.translation_text AS ai_translation_text,
            ht.translation_text AS human_translation_text,
            a.lemma,
            a.entry_number,
            COALESCE(stv.text_body, a.human_greek_text, a.greek_text, '') AS source_text
        FROM translation_runs tr
        JOIN translation_prompt_profiles p ON p.id = tr.profile_id
        JOIN translation_prompt_profile_versions pv ON pv.id = tr.profile_version_id
        JOIN paper_corpus pc ON pc.lemma_id = tr.lemma_id
        JOIN human_translations ht ON ht.id = pc.human_translation_id
        JOIN assembled_lemmas a ON a.id = tr.lemma_id
        LEFT JOIN lemma_source_text_versions stv ON stv.id = tr.source_text_version_id
        WHERE p.name = %s
          AND pv.version = %s
          AND tr.status IN ('completed', 'approved')
          AND NULLIF(BTRIM(tr.translation_text), '') IS NOT NULL
        ORDER BY pc.corpus_order, tr.run_index, tr.id
        """,
        (REPEAT_PROFILE, REPEAT_PROFILE_VERSION),
    )
    rows = [dict(row) for row in cur.fetchall()]
    if rows:
        evaluator = TranslationMetricEvaluator(("BLEU-4", "chrF++", "ROUGE-L"))
        rows = build_pair_rows(rows, metric_evaluator=evaluator)
    return rows


def fetch_model_timeline_status(cur) -> dict[str, Any]:
    cur.execute(
        f"""
        WITH {paper_kappa_review_cte_body("paper_corpus")}
        SELECT
            COUNT(DISTINCT (tr.lemma_id, tr.profile_version_id)) AS completed_pairs,
            COUNT(DISTINCT tr.lemma_id) AS completed_lemmas,
            COUNT(DISTINCT p.name) AS completed_profiles
        FROM translation_runs tr
        JOIN translation_prompt_profiles p ON p.id = tr.profile_id
        JOIN translation_prompt_profile_versions pv ON pv.id = tr.profile_version_id
        JOIN paper_corpus pc ON pc.lemma_id = tr.lemma_id
        WHERE p.name = ANY(%s)
          AND pv.version = ANY(%s)
          AND tr.status IN ('completed', 'approved')
          AND NULLIF(BTRIM(tr.translation_text), '') IS NOT NULL
        """,
        (list(MODEL_TIMELINE_PROFILES), list(MODEL_TIMELINE_PROMPT_VERSIONS)),
    )
    status = dict(cur.fetchone())
    status["expected_pairs"] = EXPECTED_MODEL_TIMELINE_PAIRS
    status["expected_profiles"] = len(MODEL_TIMELINE_PROFILES)
    return status


def fetch_repeat_queue_status(cur, repeat_rows: list[dict[str, Any]]) -> dict[str, Any]:
    cur.execute(
        f"""
        WITH {paper_kappa_review_cte_body("paper_corpus")},
        repeat_version AS (
            SELECT pv.id
            FROM translation_prompt_profiles p
            JOIN translation_prompt_profile_versions pv ON pv.profile_id = p.id
            WHERE p.name = %s
              AND pv.version = %s
            LIMIT 1
        )
        SELECT
            COUNT(*) FILTER (WHERE r.status IN ('pending', 'running')) AS open_requests,
            COALESCE(
                SUM(r.requested_runs) FILTER (WHERE r.status IN ('pending', 'running')),
                0
            ) AS open_requested_runs,
            COUNT(DISTINCT r.lemma_id)
                FILTER (WHERE r.status IN ('pending', 'running')) AS open_lemmas,
            COUNT(*) FILTER (
                WHERE r.status IN ('pending', 'running')
                  AND r.lemma_id NOT IN (SELECT lemma_id FROM paper_corpus)
            ) AS outside_paper_corpus
        FROM translation_run_requests r
        WHERE r.profile_version_id = (SELECT id FROM repeat_version)
        """,
        (REPEAT_PROFILE, REPEAT_PROFILE_VERSION),
    )
    status = dict(cur.fetchone())
    grouped = defaultdict(list)
    for row in repeat_rows:
        grouped[int(row["lemma_id"])].append(row)
    status["completed_runs"] = len(repeat_rows)
    status["completed_lemmas"] = sum(
        1 for rows in grouped.values() if len({int(row["run_index"]) for row in rows}) >= 5
    )
    status["expected_runs"] = EXPECTED_REPEAT_RUNS
    status["expected_lemmas"] = EXPECTED_PAPER_PASSAGES
    return status


def fetch_report_data() -> dict[str, Any]:
    connection = get_connection(dict_cursor=True)
    cursor = connection.cursor()
    entries = fetch_entries(cursor)
    attach_notes_and_sources(cursor, entries)
    paper_entries = fetch_paper_entries(cursor, entries)
    repeat_rows = fetch_repeatability_rows(cursor)
    timeline_status = fetch_model_timeline_status(cursor)
    queue_status = fetch_repeat_queue_status(cursor, repeat_rows)
    connection.close()

    repeats_by_lemma: dict[int, list[dict[str, Any]]] = defaultdict(list)
    for row in repeat_rows:
        repeats_by_lemma[int(row["lemma_id"])].append(row)
    for entry in paper_entries:
        entry["repeat_variants"] = repeats_by_lemma.get(int(entry["lemma_id"]), [])

    return {
        "entries": entries,
        "paper_entries": paper_entries,
        "repeatability_rows": repeat_rows,
        "model_timeline_status": timeline_status,
        "repeat_queue_status": queue_status,
    }


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, *, top: int = 80, start: int = 100, bottom: int = 80, end: int = 100) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = margins.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    for cell in row.cells:
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor.from_string(WHITE)
                run.font.bold = True
                run.font.size = Pt(8)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.size = Pt(8)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def set_document_defaults(document: Document, short_title: str) -> None:
    section = document.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")

    for style_name, size, color, before, after in (
        ("Title", 28, NAVY, 0, 12),
        ("Subtitle", 14, BLUE, 0, 12),
        ("Heading 1", 19, NAVY, 14, 6),
        ("Heading 2", 13, BLUE, 10, 4),
        ("Heading 3", 11, DARK_GREY, 7, 3),
    ):
        style = styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Kappa Metadata" not in styles:
        metadata_style = styles.add_style("Kappa Metadata", WD_STYLE_TYPE.PARAGRAPH)
    else:
        metadata_style = styles["Kappa Metadata"]
    metadata_style.font.name = "Aptos"
    metadata_style.font.size = Pt(8)
    metadata_style.font.color.rgb = RGBColor.from_string(DARK_GREY)
    metadata_style.paragraph_format.space_after = Pt(4)

    for current_section in document.sections:
        header = current_section.header.paragraphs[0]
        header.text = short_title
        header.style = styles["Kappa Metadata"]
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_page_number(current_section.footer.paragraphs[0])


def add_markdown_text(paragraph, value: Any, *, italic: bool = False) -> None:
    text = compact_whitespace(value)
    if not text:
        return
    pieces = re.split(r"(\*[^*]+\*)", text)
    for piece in pieces:
        if not piece:
            continue
        piece_is_italic = len(piece) >= 2 and piece.startswith("*") and piece.endswith("*")
        run = paragraph.add_run(piece[1:-1] if piece_is_italic else piece)
        run.italic = italic or piece_is_italic


def add_labelled_paragraph(
    document: Document,
    label: str,
    value: Any,
    *,
    label_color: str = NAVY,
    style: str | None = None,
) -> None:
    paragraph = document.add_paragraph(style=style)
    label_run = paragraph.add_run(f"{label}: ")
    label_run.bold = True
    label_run.font.color.rgb = RGBColor.from_string(label_color)
    add_markdown_text(paragraph, value)


def add_status_banner(document: Document, title: str, text: str, *, fill: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    row_properties = table.rows[0]._tr.get_or_add_trPr()
    row_properties.append(OxmlElement("w:cantSplit"))
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=140, start=160, bottom=140, end=160)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(NAVY)
    paragraph.add_run("\n")
    add_markdown_text(paragraph, text)
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def add_metric_cards(document: Document, metrics: list[tuple[str, str]]) -> None:
    table = document.add_table(rows=2, cols=len(metrics))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for index, (label, value) in enumerate(metrics):
        value_cell = table.cell(0, index)
        label_cell = table.cell(1, index)
        set_cell_shading(value_cell, PALE_BLUE)
        set_cell_shading(label_cell, PALE_GREY)
        set_cell_margins(value_cell, top=130, bottom=60)
        set_cell_margins(label_cell, top=55, bottom=90)
        value_paragraph = value_cell.paragraphs[0]
        value_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        value_run = value_paragraph.add_run(value)
        value_run.bold = True
        value_run.font.size = Pt(17)
        value_run.font.color.rgb = RGBColor.from_string(BLUE)
        label_paragraph = label_cell.paragraphs[0]
        label_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label_run = label_paragraph.add_run(label)
        label_run.font.size = Pt(7.5)
        label_run.font.color.rgb = RGBColor.from_string(DARK_GREY)


def add_cover(
    document: Document,
    *,
    title: str,
    subtitle: str,
    snapshot_date: str,
    warning_title: str,
    warning_text: str,
) -> None:
    document.add_paragraph("STEPHANOS OF BYZANTIUM", style="Kappa Metadata")
    document.add_heading(title, 0)
    subtitle_paragraph = document.add_paragraph(subtitle, style="Subtitle")
    subtitle_paragraph.paragraph_format.space_after = Pt(20)
    add_labelled_paragraph(document, "Prepared for", "Greta Hawes")
    add_labelled_paragraph(document, "Data snapshot", snapshot_date)
    add_labelled_paragraph(document, "Primary corpus", "Gabe's frozen 100-passage Kappa review tracker")
    document.add_paragraph()
    add_status_banner(
        document,
        warning_title,
        warning_text,
        fill=PALE_ORANGE,
    )
    document.add_paragraph(
        "This document separates completed evidence from work that is merely queued. "
        "Run identifiers and source choices are retained in the companion data snapshot."
    )
    document.add_page_break()


def grouped_repeat_rows(repeat_rows: list[dict[str, Any]]) -> list[tuple[int, list[dict[str, Any]]]]:
    grouped: dict[int, list[dict[str, Any]]] = defaultdict(list)
    for row in repeat_rows:
        grouped[int(row["lemma_id"])].append(row)
    return sorted(
        grouped.items(),
        key=lambda item: int(item[1][0].get("corpus_order") or item[0]),
    )


def repeat_summary(repeat_rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    result = []
    for lemma_id, rows in grouped_repeat_rows(repeat_rows):
        if len({int(row["run_index"]) for row in rows}) < 5:
            continue
        word_counts = [int(row["ai_word_count"]) for row in rows]
        result.append(
            {
                "lemma_id": lemma_id,
                "corpus_order": int(rows[0]["corpus_order"]),
                "entry_number": int(rows[0]["entry_number"]),
                "lemma": str(rows[0]["lemma"]),
                "run_count": len(rows),
                "distinct_count": len(
                    {normalize_text(str(row["ai_translation_text"])) for row in rows}
                ),
                "bleu_range": metric_range(rows, "bleu4"),
                "rouge_range": metric_range(rows, "rouge_l"),
                "chrf_range": metric_range(rows, "chrfpp"),
                "word_range": max(word_counts) - min(word_counts),
            }
        )
    return result


def create_repeatability_chart(summary: list[dict[str, Any]], output_path: Path) -> None:
    labels = [f"Κ {row['entry_number']}" for row in summary]
    x_positions = list(range(len(summary)))
    width = 0.24
    figure, axis = plt.subplots(figsize=(7.3, 3.8))
    axis.bar(
        [position - width for position in x_positions],
        [100 * row["bleu_range"] for row in summary],
        width,
        label="BLEU-4 range",
        color="#3F7EA6",
    )
    axis.bar(
        x_positions,
        [100 * row["rouge_range"] for row in summary],
        width,
        label="ROUGE-L range",
        color="#6A9F72",
    )
    axis.bar(
        [position + width for position in x_positions],
        [100 * row["chrf_range"] for row in summary],
        width,
        label="chrF++ range",
        color="#D68A3A",
    )
    axis.set_title("Variation across five repeated v3 runs", loc="left", fontweight="bold")
    axis.set_ylabel("Within-passage metric range (percentage points)")
    axis.set_xticks(x_positions, labels)
    axis.grid(axis="y", color="#D9DEE3", linewidth=0.7)
    axis.set_axisbelow(True)
    axis.spines[["top", "right"]].set_visible(False)
    axis.legend(frameon=False, ncol=3, fontsize=8, loc="upper left")
    figure.tight_layout()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    figure.savefig(output_path, dpi=220, bbox_inches="tight")
    plt.close(figure)


def add_repeatability_table(document: Document, summary: list[dict[str, Any]]) -> None:
    table = document.add_table(rows=1, cols=8)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = (
        "Passage",
        "Headword",
        "Runs",
        "Distinct",
        "BLEU range",
        "ROUGE-L range",
        "chrF++ range",
        "Word range",
    )
    for cell, label in zip(table.rows[0].cells, headers):
        cell.text = label
    set_repeat_table_header(table.rows[0])
    for row_data in summary:
        cells = table.add_row().cells
        values = (
            f"Κ {row_data['entry_number']}",
            row_data["lemma"],
            str(row_data["run_count"]),
            str(row_data["distinct_count"]),
            f"{100 * row_data['bleu_range']:.1f} pp",
            f"{100 * row_data['rouge_range']:.1f} pp",
            f"{100 * row_data['chrf_range']:.1f} pp",
            str(row_data["word_range"]),
        )
        for cell, value in zip(cells, values):
            cell.text = value
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(7.5)


def source_description(source: dict[str, Any]) -> str:
    name = compact_whitespace(source.get("english_translation"))
    greek = compact_whitespace(source.get("lemma_form") or source.get("proper_noun"))
    if name and greek and normalize_text(name) != normalize_text(greek):
        result = f"{name} ({greek})"
    else:
        result = name or greek or "Unnamed source"
    details = []
    citation = compact_whitespace(source.get("citation"))
    work = compact_whitespace(source.get("work_title"))
    qid = compact_whitespace(source.get("effective_wikidata_qid"))
    status = compact_whitespace(source.get("effective_resolution_status"))
    if citation:
        details.append(citation)
    if work:
        details.append(work)
    if qid:
        details.append(qid)
    if status:
        details.append(status.replace("_", " "))
    return f"{result} - {'; '.join(details)}" if details else result


def add_entry_notes(document: Document, entry: dict[str, Any]) -> None:
    notes = compact_whitespace(entry.get("human_notes"))
    if notes:
        add_labelled_paragraph(document, "Notes", notes)
    else:
        add_labelled_paragraph(document, "Notes", "None recorded.")

    sources = list(entry.get("cited_sources") or [])
    paragraph = document.add_paragraph()
    label = paragraph.add_run("Cited authors/sources: ")
    label.bold = True
    label.font.color.rgb = RGBColor.from_string(NAVY)
    if sources:
        paragraph.add_run(f"{len(sources)} recorded")
        for source in sources:
            item = document.add_paragraph(style="List Bullet")
            add_markdown_text(item, source_description(source))
    else:
        paragraph.add_run("None recorded.")

    commentary = list(entry.get("commentary_entries") or [])
    if commentary:
        paragraph = document.add_paragraph()
        label = paragraph.add_run("Phrase-level commentary: ")
        label.bold = True
        label.font.color.rgb = RGBColor.from_string(NAVY)
        paragraph.add_run(f"{len(commentary)} recorded")
        for comment in commentary:
            phrase = compact_whitespace(comment.get("phrase_text"))
            text = compact_whitespace(comment.get("commentary_text"))
            item = document.add_paragraph(style="List Bullet")
            if phrase:
                phrase_run = item.add_run(f"{phrase}: ")
                phrase_run.italic = True
            add_markdown_text(item, text)


def add_report_appendix_entry(document: Document, entry: dict[str, Any]) -> None:
    document.add_heading(entry["title"], level=2)
    metadata = document.add_paragraph(style="Kappa Metadata")
    metadata.add_run(
        f"Paper row {entry['corpus_order']} | Lemma ID {entry['lemma_id']} | "
        f"Greek headword {entry['lemma']} | Preferred source {entry['source_document'] or 'not recorded'}"
    )
    add_labelled_paragraph(
        document,
        "GPT-5.5 v3 production translation",
        entry.get("v3_translation") or "No v3 translation available.",
    )
    add_labelled_paragraph(
        document,
        "Current final human-reviewed translation",
        entry.get("human_translation") or "No final translation available.",
        label_color="2D6A3F",
    )
    repeat_variants = list(entry.get("repeat_variants") or [])
    if repeat_variants:
        document.add_heading("Controlled repeated v3 outputs (preliminary)", level=3)
        for row in repeat_variants:
            add_labelled_paragraph(
                document,
                f"Repeat {row['run_index']}",
                row["ai_translation_text"],
                label_color="9A5B13",
            )
    notes = compact_whitespace(entry.get("human_notes"))
    if notes:
        add_labelled_paragraph(document, "Symmachus notes", notes)


def build_report_document(
    data: dict[str, Any],
    *,
    output_path: Path,
    chart_path: Path,
    snapshot_date: str,
) -> None:
    document = Document()
    set_document_defaults(document, "Kappa translation report")
    queue = data["repeat_queue_status"]
    timeline = data["model_timeline_status"]
    repeat_rows = data["repeatability_rows"]
    summary = repeat_summary(repeat_rows)

    add_cover(
        document,
        title="Kappa Translation Report",
        subtitle="Paper corpus, model-history coverage, and preliminary v3 repeatability",
        snapshot_date=snapshot_date,
        warning_title="Repeatability variants are not finished",
        warning_text=(
            f"{queue['completed_runs']} of {queue['expected_runs']} controlled runs are complete, "
            f"covering {queue['completed_lemmas']} of {queue['expected_lemmas']} passages. "
            f"The remaining {queue['open_requests']} passage requests "
            f"({queue['open_requested_runs']} runs) are queued. Results from the completed six "
            "passages are included below but should not be treated as a corpus-wide estimate."
        ),
    )

    document.add_heading("Status at this snapshot", level=1)
    add_metric_cards(
        document,
        [
            ("Frozen paper passages", "100"),
            ("Final translations", "100"),
            (
                "Model-history pairs",
                f"{timeline['completed_pairs']:,}/{timeline['expected_pairs']:,}",
            ),
            (
                "Repeat runs",
                f"{queue['completed_runs']}/{queue['expected_runs']}",
            ),
        ],
    )
    document.add_paragraph()
    document.add_paragraph(
        "The 100-passage paper corpus is anchored to Gabe's final Kappa review tracker. "
        "For the final column, this report uses the current approved translation in Symmachus. "
        "For the AI baseline, it uses the selected GPT-5.5 prompt-version 3 run against the "
        "current preferred public Greek source."
    )

    document.add_heading("What is complete", level=1)
    complete_items = (
        "The frozen set contains exactly 100 Kappa passages, each with an approved current human translation.",
        (
            f"The core model-history grid contains {timeline['completed_pairs']:,} completed "
            f"passage/profile/prompt pairs: {timeline['expected_profiles']} model profiles, "
            "three prompt versions, and 100 passages."
        ),
        "The appendix gives the selected GPT-5.5 v3 translation followed by the current final human-reviewed version for every paper passage.",
    )
    for item in complete_items:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("What remains incomplete", level=1)
    incomplete_items = (
        (
            f"The controlled repeatability experiment has {queue['completed_runs']} of "
            f"{queue['expected_runs']} runs, with a complete five-run set for "
            f"{queue['completed_lemmas']} passages."
        ),
        (
            f"The other {queue['open_lemmas']} passages are now represented by "
            f"{queue['open_requests']} queued requests for {queue['open_requested_runs']} runs."
        ),
        "A final repeatability estimate and publication-ready comparative PDF should be regenerated only after all 500 runs are complete.",
    )
    for item in incomplete_items:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("Preliminary v3 repeatability results", level=1)
    document.add_paragraph(
        "Each completed passage was translated five times with the same GPT-5.5 v3 prompt "
        "and the same controlled profile. The ranges below compare each output with the "
        "current human translation. A larger range means that the measured similarity to "
        "the final translation varied more across the five runs."
    )
    if summary:
        create_repeatability_chart(summary, chart_path)
        document.add_picture(str(chart_path), width=Inches(6.6))
        caption = document.add_paragraph(
            "Figure 1. Within-passage metric spread for the six completed five-run sets. "
            "These passages are the first six corpus rows, not a random or representative sample."
        )
        caption.style = "Kappa Metadata"
        add_repeatability_table(document, summary)
        mean_bleu = mean(row["bleu_range"] for row in summary)
        mean_rouge = mean(row["rouge_range"] for row in summary)
        mean_chrf = mean(row["chrf_range"] for row in summary)
        distinct_sequence = ", ".join(
            f"Κ {row['entry_number']}: {row['distinct_count']}/5"
            for row in summary
        )
        document.add_paragraph(
            f"Across these six passages only, the mean within-passage range is "
            f"{100 * mean_bleu:.1f} percentage points for BLEU-4, "
            f"{100 * mean_rouge:.1f} for ROUGE-L, and "
            f"{100 * mean_chrf:.1f} for chrF++. Distinct output counts were "
            f"{distinct_sequence}. This describes the preliminary sample and is not yet a "
            "stable estimate for the full corpus."
        )
    else:
        document.add_paragraph("No repeatability runs were complete at this snapshot.")

    document.add_heading("Interpretive limits", level=1)
    limits = (
        "The six complete passages occur at the beginning of the queue and are not a random sample.",
        "Lexical overlap metrics detect surface variation; they do not, by themselves, establish philological improvement or deterioration.",
        "The production v3 translation and the five-run repeat profile are separate provenance lanes even though they use the same prompt family.",
        "The current approved Symmachus translation is treated as the final reference for this handoff. It may differ from the wording frozen in earlier tracker exports.",
    )
    for item in limits:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("Data and publication handoff", level=1)
    document.add_paragraph(
        "The companion JSON snapshot contains the exact entry, translation, run, notes, and "
        "cited-source records used for these documents. It is suitable as the basis of a later "
        "Zenodo data package after the remaining repeatability runs have completed and the "
        "final package has been checked for public-release fields."
    )
    document.add_paragraph(
        "The separate all-entry Word document supplies the editorial base requested for every "
        "official Kappa entry: formatted title, final translation where available, otherwise "
        "an UNCHECKED GPT-5.5 v3 translation, Symmachus notes, and cited authors/sources."
    )

    document.add_page_break()
    document.add_heading("Appendix: the 100 paper passages", level=1)
    document.add_paragraph(
        "For each passage, the selected GPT-5.5 v3 production translation appears first and "
        "the current final human-reviewed translation second. Where a controlled five-run set "
        "is available, all five outputs are reproduced."
    )
    for entry in data["paper_entries"]:
        add_report_appendix_entry(document, entry)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def add_all_entry(document: Document, entry: dict[str, Any]) -> None:
    document.add_heading(entry["title"], level=2)
    metadata = document.add_paragraph(style="Kappa Metadata")
    metadata.paragraph_format.keep_with_next = True
    metadata.add_run(
        f"Lemma ID {entry['lemma_id']} | Greek headword {entry['lemma']} | "
        f"Preferred source {entry['source_document'] or 'not recorded'}"
    )
    if entry["translation_status"] == "FINAL":
        add_status_banner(
            document,
            "FINAL",
            "Current approved human-reviewed translation from Symmachus.",
            fill=PALE_GREEN,
        )
        add_labelled_paragraph(
            document,
            "Translation",
            entry["human_translation"],
            label_color="2D6A3F",
        )
    else:
        add_status_banner(
            document,
            "UNCHECKED",
            "GPT-5.5 prompt-version 3 translation; not yet replaced by an approved final version.",
            fill=PALE_ORANGE,
        )
        translation = compact_whitespace(entry.get("v3_translation"))
        add_labelled_paragraph(
            document,
            "UNCHECKED - AI v3 translation",
            translation or "No GPT-5.5 v3 translation is available.",
            label_color="9A5B13",
        )
    add_entry_notes(document, entry)
    trace = document.add_paragraph(style="Kappa Metadata")
    trace.add_run(
        f"Data trace: human translation ID {entry.get('human_translation_id') or '-'}; "
        f"v3 run ID {entry.get('v3_run_id') or '-'}; "
        f"source text version ID {entry.get('source_text_version_id') or '-'}."
    )


def build_all_entries_document(
    data: dict[str, Any],
    *,
    output_path: Path,
    snapshot_date: str,
) -> None:
    document = Document()
    set_document_defaults(document, "Kappa all-entry working translation")
    final_count = sum(
        entry["translation_status"] == "FINAL" for entry in data["entries"]
    )
    unchecked_count = len(data["entries"]) - final_count
    add_cover(
        document,
        title="Kappa: All-Entry Working Translation",
        subtitle="Editorial base with Symmachus notes and cited authors/sources",
        snapshot_date=snapshot_date,
        warning_title=f"{unchecked_count} entries remain unchecked",
        warning_text=(
            f"This working copy contains {final_count} current final human-reviewed translations "
            f"and {unchecked_count} GPT-5.5 prompt-version 3 translations marked UNCHECKED. "
            "The label is repeated immediately before every non-final translation."
        ),
    )
    document.add_heading("How to use this document", level=1)
    guidance = (
        "Every official Kappa entry from Κ 1 through Κ 317 is included once.",
        "Titles follow the requested pattern: romanized headword, Billerbeck number, and Meineke reference.",
        "FINAL marks the current approved Symmachus translation.",
        "UNCHECKED marks the selected GPT-5.5 v3 translation where no final translation exists.",
        "Notes and cited authors/sources are drawn from their current Symmachus-backed records.",
        "The separately reconstructed entry without official Billerbeck/Meineke numbering is not included.",
    )
    for item in guidance:
        document.add_paragraph(item, style="List Bullet")
    document.add_page_break()
    document.add_heading("All Kappa entries", level=1)
    for entry in data["entries"]:
        add_all_entry(document, entry)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def write_data_snapshot(
    data: dict[str, Any],
    *,
    output_path: Path,
    snapshot_date: str,
) -> None:
    payload = {
        "snapshot_date": snapshot_date,
        "scope": {
            "official_kappa_entries": 317,
            "paper_corpus": "paper_kappa_review",
            "paper_passages": EXPECTED_PAPER_PASSAGES,
            "main_profile": MAIN_PROFILE,
            "main_prompt_version": MAIN_PROMPT_VERSION,
            "repeat_profile": REPEAT_PROFILE,
            "repeat_profile_version": REPEAT_PROFILE_VERSION,
            "expected_repeat_runs": EXPECTED_REPEAT_RUNS,
        },
        **data,
    }
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(
        json.dumps(json_safe(payload), ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Generate Greta's Kappa report and all-entry Word working copy."
    )
    parser.add_argument("--output-dir", type=Path, default=Path("output/doc"))
    parser.add_argument("--work-dir", type=Path, default=Path("tmp/docs/greta-kappa-report"))
    parser.add_argument("--snapshot-date", default=date.today().isoformat())
    args = parser.parse_args()

    output_dir = args.output_dir
    work_dir = args.work_dir
    output_dir.mkdir(parents=True, exist_ok=True)
    work_dir.mkdir(parents=True, exist_ok=True)
    report_path = output_dir / f"greta-kappa-translation-report-{args.snapshot_date}.docx"
    all_entries_path = output_dir / f"greta-kappa-all-entries-{args.snapshot_date}.docx"
    snapshot_path = output_dir / f"greta-kappa-data-snapshot-{args.snapshot_date}.json"
    chart_path = work_dir / "repeatability-ranges.png"

    data = fetch_report_data()
    build_report_document(
        data,
        output_path=report_path,
        chart_path=chart_path,
        snapshot_date=args.snapshot_date,
    )
    build_all_entries_document(
        data,
        output_path=all_entries_path,
        snapshot_date=args.snapshot_date,
    )
    write_data_snapshot(
        data,
        output_path=snapshot_path,
        snapshot_date=args.snapshot_date,
    )

    queue = data["repeat_queue_status"]
    timeline = data["model_timeline_status"]
    print(f"Generated report: {report_path}")
    print(f"Generated all-entry working copy: {all_entries_path}")
    print(f"Generated data snapshot: {snapshot_path}")
    print(
        "Coverage: "
        f"{len(data['paper_entries'])} paper passages; "
        f"{timeline['completed_pairs']}/{timeline['expected_pairs']} model-history pairs; "
        f"{queue['completed_runs']}/{queue['expected_runs']} repeat runs; "
        f"{queue['open_requests']} open repeat requests; "
        f"{queue['outside_paper_corpus']} outside frozen corpus."
    )


if __name__ == "__main__":
    main()
